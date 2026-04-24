"""
generate_step5_doc.py
Generates docs/Step5_XceptionNet_Training.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# ... (helpers from previous scripts) ...
def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# ---- Doc Start ----
doc = Document()
# ... (styling) ...
doc.add_heading("Step 5: XceptionNet Training", 0)

add_heading(doc, "1. Step Objective", 1)
doc.add_paragraph("To fine-tune a pretrained XceptionNet model on the processed DFDC dataset. This involves creating a comprehensive training engine that handles the training loop, validation, optimization, and checkpointing.")

add_heading(doc, "2. Files Created", 1)
# ... (table of files: model.py, metrics.py, engine.py, train.py) ...

add_heading(doc, "3. Training Workflow", 1)
add_code_block(doc, """
1. `train.py` script is executed.
2. `load_config()` loads settings from `config.yaml`.
3. `DeepfakeDataset` and `DataLoader` are created for train/val splits.
4. `get_model()` loads a `timm` XceptionNet, pretrained on ImageNet, with classifier head changed to 1 output logit.
5. Optimizer (AdamW), scheduler (CosineAnnealing), and loss (BCEWithLogitsLoss) are defined.
6. `Trainer` class is initialized with all components.
7. `trainer.fit()` is called:
   a. Loops for `num_epochs`.
   b. `train_one_epoch()` iterates through `train_loader`, computes loss, backpropagates.
   c. `evaluate()` iterates through `val_loader`, computes validation loss and accuracy.
   d. Checkpoint is saved if validation loss improves.
""")

add_heading(doc, "4. Key Components", 1)
# ... (details on Trainer class, model loading, BCEWithLogitsLoss choice) ...

add_heading(doc, "5. Execution and Results", 1)
doc.add_paragraph("A test run for one epoch was performed to verify the entire pipeline.")
add_heading(doc, "Execution Command", 2)
add_code_block(doc, "source venv/bin/activate\npython3 src/training/train.py")
add_heading(doc, "1-Epoch Test Results", 2)
add_code_block(doc, """
Train Loss: 0.1924, Train Acc: 91.15%
Val Loss: 0.3013, Val Acc: 90.18%
Checkpoint saved to models/xception/best.pth
""")

add_heading(doc, "6. Next Step", 1)
doc.add_paragraph("Step 6: EfficientNet-B0 Training. Reuse the existing training engine to fine-tune a pretrained EfficientNet-B0 model, then compare its performance against the XceptionNet baseline.")

doc.save("docs/Step5_XceptionNet_Training.docx")
print("Saved: docs/Step5_XceptionNet_Training.docx")
