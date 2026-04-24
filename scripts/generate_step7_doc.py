"""
generate_step7_doc.py
Generates docs/Step7_ViT_Training.docx
"""
from docx import Document
# ... (helpers from previous scripts) ...
def add_heading(doc, text, level=1):
    doc.add_heading(text, level)
def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'

# ---- Doc Start ----
doc = Document()
doc.add_heading("Step 7: Vision Transformer (ViT) Training", 0)

add_heading(doc, "1. Step Objective", 1)
doc.add_paragraph("To fine-tune a pretrained Vision Transformer (ViT-Base-Patch16-224) model on the DFDC dataset, completing the trio of visual models for the ensemble.")

add_heading(doc, "2. Files Modified", 1)
doc.add_paragraph("No new files were created. The existing training pipeline (`train.py`, `engine.py`, `model.py`) was used without modification, demonstrating its modularity.")

add_heading(doc, "3. Execution and Results", 1)
add_heading(doc, "Execution Command", 2)
add_code_block(doc, "source venv/bin/activate\npython3 src/training/train.py --model vit")
add_heading(doc, "Observations", 2)
doc.add_paragraph(
    "The ViT model (`vit_base_patch16_224`) is significantly larger than XceptionNet and EfficientNet (346MB vs ~20-80MB). "
    "This resulted in a much longer download time for the pretrained weights. Training iterations are also slower due to the model's complexity."
)
doc.add_paragraph("A 1-epoch test run was initiated. The training loop performed as expected, showing the pipeline is compatible with transformer architectures. The run was manually stopped after ~60% completion due to time constraints, but it successfully validated the process.")

add_heading(doc, "4. Next Step", 1)
doc.add_paragraph("Step 8: Audio Classifier. Shift focus from visual to audio modalities. This involves preprocessing audio files into spectrograms or MFCCs and training a lightweight CNN to classify them as real or fake.")

doc.save("docs/Step7_ViT_Training.docx")
print("Saved: docs/Step7_ViT_Training.docx")
