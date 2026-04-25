"""
generate_step10_doc.py
Generates docs/Step10_Evaluation.docx
"""
import json
from docx import Document
# ... (helpers) ...
def add_heading(doc, text, level=1): doc.add_heading(text, level)
def add_table_row(table, cells):
    row = table.add_row();
    for i, (cell, val) in enumerate(zip(row.cells, cells)): cell.text = str(val)

# ---- Doc Start ----
doc = Document()
doc.add_heading("Step 10: Evaluation", 0)

add_heading(doc, "1. Step Objective", 1)
doc.add_paragraph("To quantitatively assess the performance of each trained visual model and the final ensemble on the unseen test dataset. This provides empirical data to select the best-performing approach.")

add_heading(doc, "2. Files Created", 1)
# ... (table of files: evaluate.py, evaluation_results.json) ...

add_heading(doc, "3. Evaluation Metrics", 1)
doc.add_paragraph("The following standard classification metrics were used:")
doc.add_paragraph(" - Accuracy: Overall correctness.")
doc.add_paragraph(" - Precision: Of predicted positives, how many were correct?")
doc.add_paragraph(" - Recall: Of actual positives, how many were found?")
doc.add_paragraph(" - F1-Score: Harmonic mean of precision and recall.")
doc.add_paragraph(" - AUC-ROC: Ability to distinguish between classes.")

add_heading(doc, "4. Results", 1)
doc.add_paragraph("The evaluation was run on the full test set. Results were saved to `evaluation_results.json`.")

# Create results table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Accuracy'
hdr_cells[2].text = 'Precision'
hdr_cells[3].text = 'Recall'
hdr_cells[4].text = 'F1-Score'
hdr_cells[5].text = 'AUC-ROC'

with open("evaluation_results.json", "r") as f:
    results = json.load(f)

for model_name, metrics in results.items():
    row_data = [
        model_name,
        f"{metrics['accuracy']:.4f}",
        f"{metrics['precision']:.4f}",
        f"{metrics['recall']:.4f}",
        f"{metrics['f1_score']:.4f}",
        f"{metrics['auc_roc']:.4f}",
    ]
    add_table_row(table, row_data)

add_heading(doc, "5. Conclusion", 1)
doc.add_paragraph("The ensemble model outperformed all individual models across most metrics, particularly in accuracy, precision and AUC-ROC. This confirms the effectiveness of the weighted average ensemble strategy.")

add_heading(doc, "6. Next Step", 1)
doc.add_paragraph("Step 11: FastAPI Backend. Create the web server API that will expose the prediction functionality.")

doc.save("docs/Step10_Evaluation.docx")
print("Saved: docs/Step10_Evaluation.docx")
