"""
generate_step9_doc.py
Generates docs/Step9_Ensemble_Fusion.docx
"""
from docx import Document
# ... (helpers) ...
def add_heading(doc, text, level=1): doc.add_heading(text, level)
def add_code_block(doc, code):
    p = doc.add_paragraph(); p.add_run(code).font.name = 'Courier New'

# ---- Doc Start ----
doc = Document()
doc.add_heading("Step 9: Ensemble Fusion", 0)

add_heading(doc, "1. Step Objective", 1)
doc.add_paragraph("To combine the predictions from the individual visual models (XceptionNet, EfficientNet, ViT) into a single, more robust prediction. This is done by implementing different ensemble strategies.")

add_heading(doc, "2. Files Created", 1)
# ... (table of files: ensemble.py, predictor.py, verify_ensemble.py) ...

add_heading(doc, "3. Ensemble Strategy", 1)
doc.add_paragraph("The primary strategy, controlled by `config.yaml`, is a weighted average of the models' output probabilities. Weights are initially set based on assumed model performance, but can be tuned later on a validation set.")
add_code_block(doc, """
# from config.yaml
ensemble:
  strategy: "weighted_average"
  weights:
    xceptionnet: 0.35
    efficientnet: 0.30
    vit: 0.25
    audio: 0.10 # (for later)
""")

add_heading(doc, "4. Prediction Workflow", 1)
add_code_block(doc, """
1. `Predictor` class is initialized.
2. `_load_models()` loads all available model checkpoints from `models/`.
3. `predict_image(path)` is called.
4. The image is preprocessed (resized, normalized).
5. A prediction is obtained from each loaded model.
6. `ensemble.weighted_average()` combines the predictions using the configured weights.
7. A final verdict ('Real' or 'Fake') is returned based on the configured threshold.
""")

add_heading(doc, "5. Verification", 1)
doc.add_paragraph("The `scripts/verify_ensemble.py` script successfully tested the pipeline:")
add_code_block(doc, """
--- Prediction Result ---
Verdict: Real
Final Probability (Fake): 0.0017

Individual Model Probabilities (Fake):
- xceptionnet: 0.0007
- efficientnet: 0.0003
- vit: 0.0055
-------------------------
""")

add_heading(doc, "6. Next Step", 1)
doc.add_paragraph("Step 10: Evaluation. Systematically evaluate the performance of individual models and the ensemble on the test set using metrics like accuracy, precision, recall, F1-score, and AUC-ROC.")

doc.save("docs/Step9_Ensemble_Fusion.docx")
print("Saved: docs/Step9_Ensemble_Fusion.docx")
