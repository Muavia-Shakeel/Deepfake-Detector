"""
generate_step4_doc.py
Generates docs/Step4_DataLoader_Augmentation.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ... (helpers from previous scripts) ...
def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for cell, val in zip(row.cells, cells):
        cell.text = val
    return row

def shade_row(row, hex_color="2E75B6"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def white_header(row):
    shade_row(row)
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Doc Start ----
doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1.2)

# Title
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("HYBRID ENSEMBLE DEEPFAKE DETECTOR")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Step 4 Documentation: Data Loader & Augmentation")
r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Hazara University  |  {datetime.date.today().strftime('%d %B %Y')}").font.size = Pt(10)
doc.add_paragraph()

# 1. Objective
add_heading(doc, "1. Step Objective", 2, (0x1F, 0x49, 0x7D))
doc.add_paragraph("To create an efficient data loading pipeline using PyTorch's `Dataset` and `DataLoader` classes. This involves reading the master index CSV from Step 3 and applying on-the-fly data augmentation to the training set to improve model generalization and prevent overfitting.")
doc.add_paragraph()

# 2. Files Created
add_heading(doc, "2. Files Created / Modified", 2, (0x1F, 0x49, 0x7D))
t1 = doc.add_table(1, 3); t1.style = 'Table Grid'; white_header(t1.rows[0])
for c,v in zip(t1.rows[0].cells,["File", "Purpose", "Notes"]): c.text=v
files = [
    ("src/training/dataset.py", "Custom `DeepfakeDataset` class", "Reads index CSV, loads images, applies transforms."),
    ("src/training/augmentations.py", "Augmentation pipelines", "Defines train and validation transforms using Albumentations."),
    ("scripts/verify_dataset.py", "Verification script", "Loads and visualizes dataset samples to confirm correctness."),
    ("docs/step4_augmentation_verification.png", "Verification output", "Image plot showing augmented samples."),
]
for f in files: add_table_row(t1, f)
doc.add_paragraph()

# 3. Data Loading Workflow
add_heading(doc, "3. Data Loading Workflow", 2, (0x1F, 0x49, 0x7D))
add_code_block(doc, """
1. `DeepfakeDataset` is initialized for a specific split ('train', 'val', or 'test').
2. The dataset reads 'data/processed/faces_index.csv' into a pandas DataFrame.
3. It filters the DataFrame to include only rows for the requested split.
4. When `__getitem__(idx)` is called by the DataLoader:
    a. It retrieves the image path and label for the given index.
    b. Loads the image from 'data/processed/' using OpenCV.
    c. Passes the numpy image to the Albumentations transform pipeline.
    d. The pipeline applies augmentations (if training set) and normalizes the image.
    e. The transformed image (now a PyTorch tensor) and its integer label are returned.
""")
doc.add_paragraph()

# 4. Augmentation Strategy
add_heading(doc, "4. Augmentation Strategy", 2, (0x1F, 0x49, 0x7D))
doc.add_paragraph("Augmentations are applied only to the training set. The validation and test sets are only normalized to ensure consistent evaluation. The `config.yaml` file controls the parameters for each transform.")
t2 = doc.add_table(1, 3); t2.style = 'Table Grid'; white_header(t2.rows[0])
for c,v in zip(t2.rows[0].cells,["Transformation", "Parameters (from config.yaml)", "Purpose"]): c.text=v
augs = [
    ("HorizontalFlip", "p=0.5", "Encourages model to learn features invariant to facial orientation."),
    ("Rotate", "limit=10 degrees, p=0.5", "Improves robustness to slight head tilts."),
    ("RandomBrightnessContrast", "limit=0.2, p=0.5", "Simulates different lighting conditions."),
    ("GaussNoise", "var_limit=(10, 50), p=0.1", "Simulates sensor noise, improves robustness."),
    ("Normalize", "ImageNet mean/std", "Required for models pretrained on ImageNet."),
]
for a in augs: add_table_row(t2, a)
doc.add_paragraph()

# 5. Verification
add_heading(doc, "5. Verification", 2, (0x1F, 0x49, 0x7D))
doc.add_paragraph("The `scripts/verify_dataset.py` script was created to provide a visual check of the entire pipeline. It loads 8 random samples from the training set and displays the on-the-fly augmentations. The output is saved to the `docs/` folder.")
doc.add_picture('docs/step4_augmentation_verification.png', width=Inches(6.0))
doc.add_paragraph()

# 6. Key Code Snippets
add_heading(doc, "6. Key Code Snippets", 2, (0x1F, 0x49, 0x7D))
add_heading(doc, "6.1 `DeepfakeDataset.__getitem__`", 3, (0x2E, 0x75, 0xB6))
add_code_block(doc, """
def __getitem__(self, idx):
    row = self.df.iloc[idx]
    img_path = self.data_root / row['path']
    image = load_image(str(img_path))
    
    if self.transform:
        augmented = self.transform(image=image)
        image = augmented['image']
        
    label = self.class_to_idx[row['label']]
    return image, torch.tensor(label, dtype=torch.long)
""")
add_heading(doc, "6.2 `get_train_transforms`", 3, (0x2E, 0x75, 0xB6))
add_code_block(doc, """
def get_train_transforms(cfg):
    aug_cfg = cfg.augmentation.train
    return A.Compose([
        A.HorizontalFlip(p=aug_cfg.horizontal_flip),
        A.Rotate(limit=aug_cfg.rotation_limit, p=0.5),
        A.RandomBrightnessContrast(...),
        A.GaussNoise(...),
        A.Normalize(mean=IMAGENET_MEAN, std=IMAGENET_STD),
        ToTensorV2(),
    ])
""")
doc.add_paragraph()

# 7. Next Step
add_heading(doc, "7. Next Step", 2, (0x1F, 0x49, 0x7D))
doc.add_paragraph("Step 5: XceptionNet Training. With the data pipeline complete, the next step is to write the main training script. This will involve creating the training loop, defining the loss function (Cross-Entropy) and optimizer (AdamW), and fine-tuning the pretrained XceptionNet model on the DFDC dataset.")

os.makedirs("docs", exist_ok=True)
out = "docs/Step4_DataLoader_Augmentation.docx"
doc.save(out)
print(f"Saved: {out}")
