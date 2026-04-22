"""
generate_step1_doc.py
Generates docs/Step1_Environment_Setup.docx
Run: python3 scripts/generate_step1_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ---- helpers ------------------------------------------------

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, cells)):
        cell.text = val
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x1e)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # background shading via paragraph border trick (light grey)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

# ---- document -----------------------------------------------

doc = Document()

# --- Page margins
sec = doc.sections[0]
sec.top_margin    = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin   = Inches(1.2)
sec.right_margin  = Inches(1.2)

# --- Title block
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("HYBRID ENSEMBLE DEEPFAKE DETECTOR")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Step 1 Documentation: Environment & Project Setup")
run2.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    f"Hazara University, Mansehra  |  CS 7th A  |  {datetime.date.today().strftime('%d %B %Y')}"
).font.size = Pt(10)

doc.add_paragraph()

# --- Team table
add_heading(doc, "1. Team Information", level=2, color=(0x1F, 0x49, 0x7D))
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
hdr = t.rows[0]
for cell, val in zip(hdr.cells, ["S.No", "Student Name", "Roll Number"]):
    cell.text = val
    for run in cell.paragraphs[0].runs:
        run.bold = True
shade_row(hdr, "2E75B6")
for cell in hdr.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

students = [
    ("1", "Ibrahim Rabbani",   "301-221001"),
    ("2", "Taha Sohail",       "301-221051"),
    ("3", "Eisha Tur Raziya",  "301-221022"),
]
for s in students:
    add_table_row(t, s)

doc.add_paragraph()

# --- Objective
add_heading(doc, "2. Step Objective", level=2, color=(0x1F, 0x49, 0x7D))
doc.add_paragraph(
    "Step 1 establishes the complete development environment for the project. "
    "This includes creating the Python virtual environment, installing all dependencies, "
    "configuring project-wide settings via config.yaml, defining environment variables "
    "via .env, creating all required directory structures, and verifying the setup. "
    "Completing this step ensures every team member can reproduce an identical environment."
)

# --- Files created
add_heading(doc, "3. Files Created", level=2, color=(0x1F, 0x49, 0x7D))
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
hdr2 = t2.rows[0]
for cell, val in zip(hdr2.cells, ["File / Directory", "Purpose", "Notes"]):
    cell.text = val
    for run in cell.paragraphs[0].runs:
        run.bold = True
shade_row(hdr2, "2E75B6")
for cell in hdr2.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

files_data = [
    ("config.yaml",          "Central project configuration",          "All paths, model hyperparams, API settings"),
    (".env.example",         "Environment variable template",          "Copy to .env; never commit .env"),
    ("scripts/setup.sh",     "One-command environment setup script",   "Creates venv, installs deps, creates dirs"),
    ("data/raw/",            "Raw dataset storage",                    "Holds celeb-df, faceforensics, dfdc"),
    ("data/processed/",      "Preprocessed output storage",            "Frames, faces, audio subdirs"),
    ("models/*/",            "Model weight directories",               "xception, efficientnet, vit, ensemble"),
    ("notebooks/",           "Jupyter experimentation notebooks",      "Tracked with .gitkeep"),
    (".gitkeep (each dir)",  "Forces git to track empty directories",  "Removed once real files added"),
]
for row_data in files_data:
    add_table_row(t2, row_data)

doc.add_paragraph()

# --- Step-by-step execution
add_heading(doc, "4. Step-by-Step Execution", level=2, color=(0x1F, 0x49, 0x7D))

steps = [
    ("4.1 Clone Repository",
     "git clone https://github.com/Muavia-Shakeel/Deepfake-Detector.git\ncd Deepfake-Detector"),
    ("4.2 Run Setup Script",
     "bash scripts/setup.sh"),
    ("4.3 Activate Virtual Environment",
     "source venv/bin/activate"),
    ("4.4 Configure Environment Variables",
     "# Edit .env — set SECRET_KEY and DEVICE (cuda or cpu)\nnano .env"),
    ("4.5 Verify Installation",
     "python3 -c \"import torch, cv2, timm, fastapi; print('All OK')\"\npython3 -c \"import torch; print('CUDA:', torch.cuda.is_available())\""),
]

for title_text, code in steps:
    add_heading(doc, title_text, level=3, color=(0x2E, 0x75, 0xB6))
    add_code_block(doc, code)

doc.add_paragraph()

# --- config.yaml explained
add_heading(doc, "5. config.yaml — Key Sections", level=2, color=(0x1F, 0x49, 0x7D))
t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
hdr3 = t3.rows[0]
for cell, val in zip(hdr3.cells, ["Section", "Key Settings", "Rationale"]):
    cell.text = val
    for run in cell.paragraphs[0].runs:
        run.bold = True
shade_row(hdr3, "2E75B6")
for cell in hdr3.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

config_details = [
    ("paths",         "data_raw, data_processed, models, logs, uploads",         "Centralised path management; no hardcoding in source files"),
    ("datasets",      "root per dataset, train/val/test split ratios",            "70/15/15 split follows standard ML practice"),
    ("preprocessing", "frame_interval=10, face_size=224, n_mfcc=40",             "Matches ViT & EfficientNet input size; MFCC config per ASVspoof baseline"),
    ("models",        "pretrained=true, dropout, save_path per model",            "Transfer learning from ImageNet; dropout prevents overfitting"),
    ("ensemble",      "strategy=weighted_average, weights per model, threshold",  "XceptionNet highest weight; weights tuneable post-evaluation"),
    ("training",      "batch_size=32, lr=1e-4, mixed_precision=true",            "fp16 training halves VRAM usage; cosine LR schedule standard for ViT"),
    ("api",           "delete_after_inference=true, allowed_extensions",          "Privacy compliance; restricts upload types to known formats"),
]
for row_data in config_details:
    add_table_row(t3, row_data)

doc.add_paragraph()

# --- requirements
add_heading(doc, "6. Dependencies Overview", level=2, color=(0x1F, 0x49, 0x7D))
t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
hdr4 = t4.rows[0]
for cell, val in zip(hdr4.cells, ["Package", "Version", "Purpose"]):
    cell.text = val
    for run in cell.paragraphs[0].runs:
        run.bold = True
shade_row(hdr4, "2E75B6")
for cell in hdr4.cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

deps = [
    ("torch / torchvision", ">=2.2.0",  "Core deep learning framework"),
    ("timm",                ">=0.9.12", "Pretrained XceptionNet, EfficientNet, ViT models"),
    ("opencv-python",       ">=4.9.0",  "Frame extraction from video"),
    ("mtcnn",               ">=0.1.1",  "Multi-task CNN face detection"),
    ("librosa",             ">=0.10.1", "Audio feature extraction (MFCC, spectrogram)"),
    ("fastapi + uvicorn",   ">=0.110.0","Async REST API backend"),
    ("albumentations",      ">=1.3.1",  "GPU-accelerated image augmentations"),
    ("grad-cam",            ">=1.4.8",  "Grad-CAM explainability heatmaps"),
    ("scikit-learn",        ">=1.4.0",  "Metrics, ensemble stacking"),
]
for row_data in deps:
    add_table_row(t4, row_data)

doc.add_paragraph()

# --- Directory tree
add_heading(doc, "7. Final Directory Structure After Step 1", level=2, color=(0x1F, 0x49, 0x7D))
tree = """\
fyp-deepfake/
├── config.yaml               ← central config
├── .env.example              ← env template
├── .env                      ← local env (not committed)
├── .gitignore
├── README.md
├── requirements.txt
├── scripts/
│   └── setup.sh              ← environment setup script
├── src/
│   ├── __init__.py
│   ├── api/__init__.py
│   ├── inference/__init__.py
│   ├── preprocessing/__init__.py
│   ├── training/__init__.py
│   └── utils/__init__.py
├── tests/__init__.py
├── notebooks/                ← .gitkeep
├── data/
│   ├── raw/
│   │   ├── celeb-df/         ← .gitkeep (datasets not committed)
│   │   ├── faceforensics/    ← .gitkeep
│   │   └── dfdc/             ← .gitkeep
│   └── processed/
│       ├── frames/           ← .gitkeep
│       ├── faces/            ← .gitkeep
│       └── audio/            ← .gitkeep
└── models/
    ├── xception/             ← .gitkeep
    ├── efficientnet/         ← .gitkeep
    ├── vit/                  ← .gitkeep
    └── ensemble/             ← .gitkeep"""
add_code_block(doc, tree)

doc.add_paragraph()

# --- Verification checklist
add_heading(doc, "8. Verification Checklist", level=2, color=(0x1F, 0x49, 0x7D))
checks = [
    "Python >= 3.10 confirmed",
    "venv created at project root",
    "All packages in requirements.txt installed without error",
    "config.yaml present and parseable (pyyaml)",
    ".env created from .env.example",
    "CUDA available (if GPU machine): torch.cuda.is_available() == True",
    "All data/ and models/ subdirectories created",
    "git remote origin set to Muavia-Shakeel/Deepfake-Detector",
    "Step 1 commit pushed to main branch on GitHub",
]
for c in checks:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)

doc.add_paragraph()

# --- Next step
add_heading(doc, "9. Next Step", level=2, color=(0x1F, 0x49, 0x7D))
doc.add_paragraph(
    "Step 2: Dataset Download & Inspection — download Celeb-DF v2, FaceForensics++ (c23), "
    "and a DFDC subset; inspect sample counts, verify real/fake ratios, generate dataset "
    "statistics report."
)

# --- footer
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    "Hybrid Ensemble Deepfake Detector  ·  Hazara University FYP  ·  Step 1 of 15"
)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

# ---- save -------------------------------------------------
os.makedirs("docs", exist_ok=True)
out = "docs/Step1_Environment_Setup.docx"
doc.save(out)
print(f"Saved: {out}")
