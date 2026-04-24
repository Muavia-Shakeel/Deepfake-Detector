"""
generate_step3_doc.py
Generates docs/Step3_Preprocessing_Pipeline.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ... (helpers from previous scripts) ...
def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for cell, val in zip(row.cells, cells):
        cell.text = val
    return row

def shade_row(row, hex_color="2E75B6"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def white_header(row):
    shade_row(row)
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Doc Start ----
doc = Document()
# ... (styling from previous) ...
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1.2)

# Title
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("HYBRID ENSEMBLE DEEPFAKE DETECTOR")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Step 3 Documentation: Preprocessing Pipeline")
r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Hazara University  |  {datetime.date.today().strftime('%d %B %Y')}").font.size = Pt(10)
doc.add_paragraph()

# 1. Objective
add_heading(doc, "1. Step Objective", 2, (0x1F, 0x49, 0x7D))
doc.add_paragraph("Create a robust, reusable pipeline to process raw image data into a format suitable for deep learning model training. This involves resizing, normalization, and creating a master index file for easy data loading.")
doc.add_paragraph()

# 2. Files Created
add_heading(doc, "2. Files Created / Modified", 2, (0x1F, 0x49, 0x7D))
t1 = doc.add_table(1, 3); t1.style = 'Table Grid'; white_header(t1.rows[0])
for c,v in zip(t1.rows[0].cells,["File", "Purpose", "Notes"]): c.text=v
files = [
    ("src/preprocessing/preprocessor.py", "Core image processing functions", "Resizes to 224x224, normalizes (ImageNet), saves PNGs."),
    ("src/preprocessing/dataset_index.py", "Builds master CSV index", "Scans processed dir, maps files to split/label."),
    ("src/utils/config_loader.py", "Loads config.yaml", "Provides dot-notation access to project config."),
    ("data/processed/faces/", "Output directory for processed images", "Mirrors raw data structure (train/val/test/real/fake)."),
    ("data/processed/faces_index.csv", "Master index file", "Used by PyTorch DataLoader in next step."),
]
for f in files: add_table_row(t1, f)
doc.add_paragraph()

# 3. Preprocessing Workflow
add_heading(doc, "3. Preprocessing Workflow", 2, (0x1F, 0x49, 0x7D))
add_code_block(doc, """
1. Load config from 'config.yaml'
2. Scan 'data/raw/dfdc/dfdc_frames/' for all .png files
3. For each image:
    a. Load image with OpenCV (BGR -> RGB)
    b. Center-crop to square (preserve aspect ratio)
    c. Resize to 224x224 (Lanczos interpolation)
    d. Save as uint8 PNG to 'data/processed/faces/'
       (maintaining split/label subdirectory structure)
4. After all images processed, scan 'data/processed/faces/'
5. Generate 'faces_index.csv' with columns: split, label, path
""")
doc.add_paragraph()

# 4. Key Functions
add_heading(doc, "4. Key Functions", 2, (0x1F, 0x49, 0x7D))
t2 = doc.add_table(1, 3); t2.style = 'Table Grid'; white_header(t2.rows[0])
for c,v in zip(t2.rows[0].cells,["Function", "File", "Purpose"]): c.text=v
funcs = [
    ("process_directory()", "preprocessor.py", "Main batch processing loop. Iterates all raw images."),
    ("preprocess_image()", "preprocessor.py", "Loads, resizes, and optionally normalizes a single image."),
    ("build_index()", "dataset_index.py", "Scans output and generates the final CSV index."),
    ("load_config()", "config_loader.py", "Loads and parses the central config.yaml file."),
]
for f in funcs: add_table_row(t2, f)
doc.add_paragraph()

# 5. Execution
add_heading(doc, "5. Execution & Results", 2, (0x1F, 0x49, 0x7D))
doc.add_paragraph("The pipeline was executed on the DFDC dataset. Initial run failed due to missing `opencv-python` dependency, which was resolved by installing it into the virtual environment.")
add_heading(doc, "5.1 Execution Commands", 3, (0x2E, 0x75, 0xB6))
add_code_block(doc, """
# Activate venv
source venv/bin/activate

# Run main processing script (takes time)
python3 -c "from src.preprocessing import preprocessor; preprocessor.process_directory('data/raw/dfdc/dfdc_frames', 'data/processed/faces')"

# Build the index CSV
python3 src/preprocessing/dataset_index.py
""")
add_heading(doc, "5.2 Output Verification", 3, (0x2E, 0x75, 0xB6))
t3 = doc.add_table(1, 2); t3.style = 'Table Grid'; white_header(t3.rows[0])
for c,v in zip(t3.rows[0].cells,["Verification Step", "Result"]): c.text=v
checks = [
    ("Count processed images", "3,905 files found in data/processed/faces/"),
    ("Check index file rows", "faces_index.csv contains 3,905 data rows + header"),
    ("Check image dimensions", "All output images are 224x224 pixels"),
]
for ch in checks: add_table_row(t3, ch)
doc.add_paragraph()

# 6. Next Step
add_heading(doc, "6. Next Step", 2, (0x1F, 0x49, 0x7D))
doc.add_paragraph("Step 4: Data Loader & Augmentation. Create a custom PyTorch `Dataset` class that reads `faces_index.csv`. Implement data augmentation transforms (flips, rotations) for the training set using Albumentations.")

# ... (Footer) ...
os.makedirs("docs", exist_ok=True)
out = "docs/Step3_Preprocessing_Pipeline.docx"
doc.save(out)
print(f"Saved: {out}")
