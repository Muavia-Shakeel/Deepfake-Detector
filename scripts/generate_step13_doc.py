"""
generate_step13_doc.py
Generates docs/Step13_Docker_Deployment.docx
"""
from docx import Document
# ... (helpers) ...
def add_heading(doc, text, level=1): doc.add_heading(text, level)
def add_code_block(doc, code):
    p = doc.add_paragraph(); p.add_run(code).font.name = 'Courier New'

# ---- Doc Start ----
doc = Document()
doc.add_heading("Step 13: Docker & Deployment", 0)

add_heading(doc, "1. Step Objective", 1)
doc.add_paragraph("To containerize the entire application using Docker, ensuring a portable and reproducible environment for both development and deployment.")

add_heading(doc, "2. Files Created", 1)
# ... (table of files: Dockerfile, docker-compose.yml, .dockerignore) ...

add_heading(doc, "3. Dockerfile Strategy", 1)
doc.add_paragraph("A multi-stage build is used to keep the final image size small:")
add_code_block(doc, """
- Stage 1 ('builder'): Installs build tools and all Python dependencies into a virtual environment.
- Stage 2 ('final'): Copies the virtual environment from the builder stage and the application code. This means the final image doesn't contain build-essential and other temporary build packages.
- Model Caching: A `RUN` command is used to trigger the download of all models during the build process, so they are cached in a Docker layer.
""")

add_heading(doc, "4. Docker Compose", 1)
doc.add_paragraph("`docker-compose.yml` is configured to build and run the service. It maps port 8000 and includes a volume for the `uploads` directory to persist user uploads if needed. It also includes commented-out GPU support for production deployments.")

add_heading(doc, "5. Usage", 1)
add_heading(doc, "Build the image", 2)
add_code_block(doc, "docker compose build")
add_heading(doc, "Run the container", 2)
add_code_block(doc, "docker compose up")
doc.add_paragraph("The application will then be available at http://localhost:8000")

add_heading(doc, "6. Next Step", 1)
doc.add_paragraph("The project is now fully containerized and ready for deployment to any platform that supports Docker. All core technical requirements from the proposal have been met. Final documentation and project cleanup are the remaining tasks.")

doc.save("docs/Step13_Docker_Deployment.docx")
print("Saved: docs/Step13_Docker_Deployment.docx")
