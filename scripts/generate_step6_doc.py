"""
generate_step6_doc.py
Generates docs/Step6_EfficientNet_Training.docx
"""
from docx import Document
# ... (helpers from previous scripts) ...
def add_heading(doc, text, level=1):
    doc.add_heading(text, level)
def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'

# ---- Doc Start ----
doc = Document()
doc.add_heading("Step 6: EfficientNet-B0 Training", 0)

add_heading(doc, "1. Step Objective", 1)
doc.add_paragraph("To reuse and validate the training engine by fine-tuning a second, different architecture (EfficientNet-B0). This step demonstrates the modularity of the training pipeline and provides a second model for the final ensemble.")

add_heading(doc, "2. Files Modified", 1)
# ... (table of files: train.py, engine.py, model.py) ...

add_heading(doc, "3. Key Changes", 1)
doc.add_paragraph("The core training logic in `engine.py` remained unchanged. Key modifications to support multi-model training include:")
add_code_block(doc, """
- `train.py`: Added `argparse` to accept a `--model` command-line argument.
- `train.py`: Passes the model-specific config section (e.g., `cfg.models.efficientnet`) to the Trainer.
- `engine.py`: The `Trainer` now uses the passed `model_cfg` to determine the correct checkpoint save path.
- `model.py`: `get_model()` now dynamically handles 'xceptionnet', 'efficientnet', and 'vit' variants.
""")

add_heading(doc, "4. Execution and Results", 1)
add_heading(doc, "Execution Command", 2)
add_code_block(doc, "source venv/bin/activate\npython3 src/training/train.py --model efficientnet")
add_heading(doc, "Training Snippet (6 Epochs)", 2)
add_code_block(doc, """
Epoch 1/30: Train Acc: 86.66%, Val Acc: 82.08%
Epoch 2/30: Train Acc: 92.95%, Val Acc: 85.22%
Epoch 3/30: Train Acc: 94.02%, Val Acc: 81.09%
Epoch 4/30: Train Acc: 95.05%, Val Acc: 86.61%
Epoch 5/30: Train Acc: 95.51%, Val Acc: 85.13%
Epoch 6/30: Train Acc: 96.35%, Val Acc: 87.38%
... (timeout)
""")
doc.add_paragraph("Best validation accuracy after 6 epochs was 87.38%. Checkpoint saved to `models/efficientnet/best.pth`.")

add_heading(doc, "5. Next Step", 1)
doc.add_paragraph("Step 7: Vision Transformer (ViT) Training. Train the third and final visual model using the same training pipeline.")

doc.save("docs/Step6_EfficientNet_Training.docx")
print("Saved: docs/Step6_EfficientNet_Training.docx")
