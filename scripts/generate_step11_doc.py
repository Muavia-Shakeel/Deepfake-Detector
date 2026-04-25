"""
generate_step11_doc.py
Generates docs/Step11_FastAPI_Backend.docx
"""
from docx import Document
# ... (helpers) ...
def add_heading(doc, text, level=1): doc.add_heading(text, level)
def add_code_block(doc, code):
    p = doc.add_paragraph(); p.add_run(code).font.name = 'Courier New'

# ---- Doc Start ----
doc = Document()
doc.add_heading("Step 11: FastAPI Backend", 0)

add_heading(doc, "1. Step Objective", 1)
doc.add_paragraph("To create a web-based API that exposes the deepfake detection functionality, allowing users to upload media for analysis.")

add_heading(doc, "2. Files Created", 1)
# ... (table of files: main.py, video_utils.py) ...

add_heading(doc, "3. API Endpoint", 1)
add_heading(doc, "POST /predict", 2)
doc.add_paragraph("Accepts a multipart/form-data request with a single file upload.")
doc.add_paragraph("- For images: returns a prediction immediately.")
doc.add_paragraph("- For videos: extracts frames, runs prediction on each, and returns the average probability.")

add_heading(doc, "4. Execution and Testing", 1)
add_heading(doc, "Running the Server", 2)
add_code_block(doc, "source venv/bin/activate\npython3 src/api/main.py")
add_heading(doc, "Testing with cURL", 2)
add_code_block(doc, "curl -X POST -F \"file=@<path_to_image>\" http://127.0.0.1:8000/predict")
add_heading(doc, "Sample Response", 2)
add_code_block(doc, """
{
  "final_probability": 0.2403,
  "verdict": "Real",
  "individual_probabilities": {
    "xceptionnet": 0.2209,
    "efficientnet": 0.4757,
    "vit": 0.0
  }
}
""")

add_heading(doc, "5. Next Step", 1)
doc.add_paragraph("Step 12: Frontend. Build a user interface to interact with the API.")

doc.save("docs/Step11_FastAPI_Backend.docx")
print("Saved: docs/Step11_FastAPI_Backend.docx")
