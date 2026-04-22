"""
generate_step2_doc.py
Generates docs/Step2_Dataset_Download_Inspection.docx
Run: python3 scripts/generate_step2_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ---- helpers (same as step1) --------------------------------

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, cells)):
        cell.text = val
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

def shade_row(row, hex_color="2E75B6"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x1e)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def white_header(row):
    shade_row(row)
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- document -----------------------------------------------

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1.2)

# Title
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("HYBRID ENSEMBLE DEEPFAKE DETECTOR")
r.bold = True; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Step 2 Documentation: Dataset Download & Inspection")
r2.bold = True; r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(
    f"Hazara University, Mansehra  |  CS 7th A  |  {datetime.date.today().strftime('%d %B %Y')}"
).font.size = Pt(10)

doc.add_paragraph()

# 1. Objective
add_heading(doc, "1. Step Objective", level=2, color=(0x1F, 0x49, 0x7D))
doc.add_paragraph(
    "Step 2 covers acquisition and initial inspection of training datasets. "
    "The goal is to ensure sufficient and balanced data is available before any preprocessing "
    "or model training begins. The inspection phase produces statistics on class distribution, "
    "split ratios, and image properties to guide later preprocessing decisions."
)
doc.add_paragraph()

# 2. Datasets Used
add_heading(doc, "2. Datasets Acquired", level=2, color=(0x1F, 0x49, 0x7D))
t1 = doc.add_table(rows=1, cols=4)
t1.style = 'Table Grid'
white_header(t1.rows[0])
for cell, val in zip(t1.rows[0].cells,
                     ["Dataset", "Source", "Status", "Location"]):
    cell.text = val

rows = [
    ("DFDC (Deepfake Detection Challenge)",
     "Kaggle — archive.zip (pre-extracted frames)",
     "✅ Fully extracted — 50,000 face frames",
     "data/raw/dfdc/dfdc_frames/"),
    ("FaceForensics++ c23",
     "Direct download (browser)",
     "⚠️ Partial — 1.3 GB downloaded, incomplete",
     "archive (1).zip.crdownload"),
    ("Celeb-DF v2",
     "Author request form",
     "🔜 Pending — request to be submitted",
     "data/raw/celeb-df/ (empty)"),
    ("ASVspoof (audio)",
     "ASVspoof challenge website",
     "🔜 Pending — Step 8 (audio classifier)",
     "data/raw/ (future)"),
]
for r in rows:
    add_table_row(t1, r)
doc.add_paragraph()

# 3. DFDC Statistics
add_heading(doc, "3. DFDC Dataset — Statistics", level=2, color=(0x1F, 0x49, 0x7D))
doc.add_paragraph(
    "The DFDC archive contained pre-extracted, pre-labeled face frame crops (PNG), "
    "already split into train/val/test subsets. No additional split needed. "
    "CSV log (dfdc_log.csv) indexes all 49,029 valid frames."
)

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
white_header(t2.rows[0])
for cell, val in zip(t2.rows[0].cells, ["Split", "Real", "Fake", "Total"]):
    cell.text = val

stats = [
    ("Train", "18,000", "18,000", "36,000"),
    ("Val",   "3,500",  "3,500",  "7,000"),
    ("Test",  "2,529",  "3,500",  "6,029"),
    ("Total", "24,029", "25,000", "49,029"),
]
for r in stats:
    add_table_row(t2, r)

doc.add_paragraph()

add_heading(doc, "3.1 Class Balance Analysis", level=3, color=(0x2E, 0x75, 0xB6))
doc.add_paragraph(
    "Train split: perfectly balanced (50% real / 50% fake) — ideal for unbiased training. "
    "Test split: slight fake imbalance (58% fake / 42% real) — reflects real-world deployment "
    "conditions where fake content is more prevalent. No resampling required for train/val; "
    "class weights may optionally be applied during training given test imbalance."
)
doc.add_paragraph()

add_heading(doc, "3.2 Image Properties", level=3, color=(0x2E, 0x75, 0xB6))
t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Table Grid'
white_header(t3.rows[0])
for cell, val in zip(t3.rows[0].cells, ["Property", "Value"]):
    cell.text = val
props = [
    ("Format",          "PNG (lossless)"),
    ("Content",         "Cropped face regions"),
    ("Naming pattern",  "<video_id>_<frame_number>_<face_id>.png"),
    ("Source labels",   "Provided in dfdc_log.csv (label, label_encoded columns)"),
    ("Pre-split",       "Yes — train/val/test folders present"),
    ("Resize needed",   "Yes — target 224×224 for model input (done in preprocessing step)"),
]
for r in props:
    add_table_row(t3, r)
doc.add_paragraph()

# 4. FaceForensics++
add_heading(doc, "4. FaceForensics++ — Status", level=2, color=(0x1F, 0x49, 0x7D))
doc.add_paragraph(
    "A partial FaceForensics++ c23 archive (1.3 GB) exists as 'archive (1).zip.crdownload'. "
    "The file header confirms it is a valid ZIP beginning with FaceForensics++_C23/DeepFakeDetect content. "
    "Download was interrupted mid-transfer. Action required: resume or re-download from source."
)
add_heading(doc, "4.1 Resume Download Instructions", level=3, color=(0x2E, 0x75, 0xB6))
doc.add_paragraph(
    "FaceForensics++ requires a data access request. Submit the form at:"
)
add_code_block(doc, "https://docs.google.com/forms/d/e/1FAIpQLSdRRR3L5zAv6tQ_CKxmK4W96Os0j_oQHpBiW8ctQfxBl8Kfsg/viewform")
doc.add_paragraph(
    "Once credentials received, use the official download script (download-FaceForensics.py) "
    "with compression level c23 (medium quality — standard for research):"
)
add_code_block(doc, "python download-FaceForensics.py data/raw/faceforensics -d all -c c23 -t videos")
doc.add_paragraph()

# 5. Celeb-DF v2
add_heading(doc, "5. Celeb-DF v2 — Acquisition Plan", level=2, color=(0x1F, 0x49, 0x7D))
doc.add_paragraph(
    "Celeb-DF v2 is hosted by authors at SUNY Albany. Access via:"
)
add_code_block(doc, "https://github.com/yuezunli/celeb-deepfakeforensics  (README has Google Form link)")
t4 = doc.add_table(rows=1, cols=2)
t4.style = 'Table Grid'
white_header(t4.rows[0])
for cell, val in zip(t4.rows[0].cells, ["Property", "Value"]):
    cell.text = val
celeb_props = [
    ("Real videos",  "590"),
    ("Fake videos",  "5,639"),
    ("Approx size",  "~2 GB"),
    ("Format",       "MP4"),
    ("Access",       "Google Form → Google Drive link"),
]
for r in celeb_props:
    add_table_row(t4, r)
doc.add_paragraph()

# 6. Data directory structure
add_heading(doc, "6. Data Directory Structure After Step 2", level=2, color=(0x1F, 0x49, 0x7D))
add_code_block(doc, """\
data/
├── raw/
│   ├── dfdc/
│   │   ├── dfdc_frames/
│   │   │   ├── train/
│   │   │   │   ├── real/   ← 18,000 PNG face frames
│   │   │   │   └── fake/   ← 18,000 PNG face frames
│   │   │   ├── val/
│   │   │   │   ├── real/   ← 3,500 frames
│   │   │   │   └── fake/   ← 3,500 frames
│   │   │   └── test/
│   │   │       ├── real/   ← 2,529 frames
│   │   │       └── fake/   ← 3,500 frames
│   │   └── dfdc_log.csv    ← index with labels + split info
│   ├── faceforensics/       ← PENDING (partial download)
│   └── celeb-df/            ← PENDING (access request)
└── processed/               ← filled in Step 3 (preprocessing)""")
doc.add_paragraph()

# 7. Inspection Commands
add_heading(doc, "7. Inspection Commands Used", level=2, color=(0x1F, 0x49, 0x7D))
cmds = [
    ("Count extracted files",
     "find data/raw/dfdc -type f | wc -l"),
    ("Inspect archive contents",
     "unzip -l archive.zip | head -20"),
    ("Dataset split statistics",
     "python3 -c \"import pandas as pd; df=pd.read_csv('data/raw/dfdc/dfdc_log.csv'); print(df.groupby(['split','label']).size())\""),
    ("Check partial download type",
     "python3 -c \"f=open('archive (1).zip.crdownload','rb'); print(f.read(64))\""),
]
t5 = doc.add_table(rows=1, cols=2)
t5.style = 'Table Grid'
white_header(t5.rows[0])
for cell, val in zip(t5.rows[0].cells, ["Purpose", "Command"]):
    cell.text = val
for r in cmds:
    add_table_row(t5, r)
doc.add_paragraph()

# 8. Key Observations
add_heading(doc, "8. Key Observations & Decisions", level=2, color=(0x1F, 0x49, 0x7D))
obs = [
    "DFDC dataset arrives pre-extracted as face crops — skips raw video frame extraction for this source.",
    "DFDC frames still require resize to 224×224 and normalization before model input.",
    "dfdc_log.csv provides ground-truth labels — use it as primary index for DataLoader in Step 4.",
    "FaceForensics++ and Celeb-DF will add diversity to training set (different manipulation methods).",
    "Start model development on DFDC alone — sufficient size (36k train) for initial training.",
    "class_weight can be used in loss function if test imbalance affects initial results.",
]
for o in obs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(o)
doc.add_paragraph()

# 9. Verification Checklist
add_heading(doc, "9. Verification Checklist", level=2, color=(0x1F, 0x49, 0x7D))
checks = [
    "archive.zip extracted to data/raw/dfdc/ (50,001 files)",
    "dfdc_log.csv readable and contains 49,029 rows",
    "Train/val/test folder structure confirmed under dfdc_frames/",
    "Real/fake subdirectories present in each split",
    "Image format verified as PNG face crops",
    "Class balance confirmed: train 50/50, val 50/50, test 42/58",
    "kaggle.json secured in ~/.kaggle/ only (not in project folder)",
    "FaceForensics++ access form noted for submission",
    "Celeb-DF v2 acquisition plan documented",
]
for c in checks:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)
doc.add_paragraph()

# 10. Next Step
add_heading(doc, "10. Next Step", level=2, color=(0x1F, 0x49, 0x7D))
doc.add_paragraph(
    "Step 3: Preprocessing Pipeline — build src/preprocessing/ module to: "
    "(1) resize face frames to 224×224, "
    "(2) normalize pixel values per ImageNet stats, "
    "(3) apply face alignment where needed, "
    "(4) generate processed dataset index CSV for DataLoader use."
)

# Footer
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Hybrid Ensemble Deepfake Detector  ·  Hazara University FYP  ·  Step 2 of 15")
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

os.makedirs("docs", exist_ok=True)
out = "docs/Step2_Dataset_Download_Inspection.docx"
doc.save(out)
print(f"Saved: {out}")
